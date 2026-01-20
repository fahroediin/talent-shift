"""
Script to generate sample CV PDF for testing
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_sample_cv():
    doc = Document()
    
    # Title
    title = doc.add_heading('ANDI WIJAYA', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Contact Info
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run('Email: andi.wijaya@email.com | Phone: +62 812-3456-7890 | Location: Jakarta')
    
    doc.add_paragraph()
    
    # Summary
    doc.add_heading('SUMMARY', level=1)
    doc.add_paragraph(
        'Experienced Backend Developer with 5 years of experience in building scalable '
        'web applications using Python, SQL, and RESTful APIs. Passionate about clean '
        'code and best practices.'
    )
    
    # Education
    doc.add_heading('EDUCATION', level=1)
    edu = doc.add_paragraph()
    edu.add_run('S1 Teknik Informatika').bold = True
    edu.add_run('\nUniversitas Indonesia (UI)')
    edu.add_run('\nGraduated: 2019 | GPA: 3.75/4.00')
    
    # Experience
    doc.add_heading('WORK EXPERIENCE', level=1)
    
    exp1 = doc.add_paragraph()
    exp1.add_run('Senior Backend Developer').bold = True
    exp1.add_run(' - PT Gojek Indonesia')
    exp1.add_run('\nJanuary 2022 - Present (3 years)')
    
    doc.add_paragraph('• Developed microservices using Python and FastAPI', style='List Bullet')
    doc.add_paragraph('• Designed and optimized SQL databases using PostgreSQL', style='List Bullet')
    doc.add_paragraph('• Built REST API endpoints serving 1M+ requests/day', style='List Bullet')
    doc.add_paragraph('• Implemented Docker containerization for deployments', style='List Bullet')
    doc.add_paragraph('• Managed services on AWS (EC2, RDS, S3)', style='List Bullet')
    
    exp2 = doc.add_paragraph()
    exp2.add_run('Backend Developer').bold = True
    exp2.add_run(' - PT Tokopedia')
    exp2.add_run('\nJuly 2019 - December 2021 (2.5 years)')
    
    doc.add_paragraph('• Built backend services using Python Django framework', style='List Bullet')
    doc.add_paragraph('• Wrote complex SQL queries for data analytics', style='List Bullet')
    doc.add_paragraph('• Developed REST API for mobile applications', style='List Bullet')
    
    # Skills
    doc.add_heading('TECHNICAL SKILLS', level=1)
    doc.add_paragraph('Programming Languages: Python, JavaScript, Go')
    doc.add_paragraph('Databases: PostgreSQL, MySQL, MongoDB, Redis')
    doc.add_paragraph('Frameworks: FastAPI, Django, Flask')
    doc.add_paragraph('DevOps: Docker, Kubernetes, AWS, GCP')
    doc.add_paragraph('Other: REST API, GraphQL, Microservices, Linux, Git')
    
    # Certifications
    doc.add_heading('CERTIFICATIONS & TRAINING', level=1)
    doc.add_paragraph('• Hacktiv8 Full Stack Developer Program (2019)', style='List Bullet')
    doc.add_paragraph('• AWS Certified Developer Associate (2021)', style='List Bullet')
    doc.add_paragraph('• Dicoding - Backend Developer Path (2020)', style='List Bullet')
    
    # Portfolio
    doc.add_heading('PORTFOLIO', level=1)
    doc.add_paragraph('GitHub: https://github.com/andiwijaya')
    doc.add_paragraph('LinkedIn: https://linkedin.com/in/andiwijaya')
    doc.add_paragraph('Personal Website: https://andiwijaya.dev')
    
    # Save
    doc.save('samples/sample_cv_andi_wijaya.docx')
    print("CV created: samples/sample_cv_andi_wijaya.docx")

if __name__ == "__main__":
    create_sample_cv()
