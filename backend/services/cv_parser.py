"""
CV Parser Service
Extracts structured data from PDF and DOCX files
"""

import re
import fitz  # PyMuPDF
from docx import Document
from io import BytesIO
from typing import Optional, List
from models.schemas import ParsedCV


class CVParser:
    """
    CV Parser that extracts structured information from CV files
    Supports PDF and DOCX formats
    """
    
    # Regex patterns for extraction
    EMAIL_PATTERN = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    PHONE_PATTERN = r'(?:\+62|62|0)[\s-]?(?:\d{2,4})[\s-]?(?:\d{3,4})[\s-]?(?:\d{3,4})'
    
    # URL patterns
    GITHUB_PATTERN = r'(?:https?://)?(?:www\.)?github\.com/[\w-]+'
    LINKEDIN_PATTERN = r'(?:https?://)?(?:www\.)?linkedin\.com/in/[\w-]+'
    WEBSITE_PATTERN = r'https?://[\w.-]+\.[a-zA-Z]{2,}(?:/[\w./-]*)?'
    
    # Education level keywords - use specific patterns to avoid false positives
    EDUCATION_LEVELS = {
        'S3': ['doktor', 'phd', 'doctoral', 'doctorate', 'doctor of'],
        'S2': ['magister', 'master', 'msc', 'mba', 'mm', 'mt', 'mti', "master's"],
        'S1': ['sarjana', 'bachelor', 'skom', 's.kom', 's.t', 's.e', "bachelor's"],
        'D3': ['diploma', 'ahli madya', 'd3', 'd-3'],
        'D4': ['sarjana terapan', 'd4', 'd-4'],
        'SMA': ['sma', 'smk', 'high school', 'smu', 'slta'],
    }
    
    # Education level patterns that need word boundaries (avoid AWS S3 etc)
    EDUCATION_STANDALONE = {
        'S3': r'\bs3\b(?!\s*(?:bucket|storage|aws|amazon))',  # S3 but not AWS S3
        'S2': r'\bs2\b',
        'S1': r'\bs1\b',
    }
    
    # Known bootcamp/training providers
    BOOTCAMP_PROVIDERS = [
        'hacktiv8', 'binar', 'binar academy', 'dicoding', 'purwadhika', 
        'sanbercode', 'glints academy', 'coursera', 'udemy', 'linkedin learning',
        'google career', 'aws training', 'meta blueprint'
    ]
    
    # Common skill keywords
    SKILL_KEYWORDS = [
        # Programming languages
        'python', 'javascript', 'typescript', 'java', 'c++', 'c#', 'go', 'golang',
        'ruby', 'php', 'swift', 'kotlin', 'rust', 'scala', 'r',
        # Web frameworks
        'react', 'vue', 'angular', 'nextjs', 'next.js', 'express', 'fastapi',
        'django', 'flask', 'spring', 'laravel', 'rails',
        # Low-code platforms
        'outsystems', 'mendix', 'appian', 'power apps',
        # Databases
        'sql', 'mysql', 'postgresql', 'postgres', 'mongodb', 'redis', 'elasticsearch',
        'oracle', 'sqlite', 'dynamodb',
        # DevOps & Cloud
        'docker', 'kubernetes', 'k8s', 'aws', 'gcp', 'azure', 'terraform',
        'jenkins', 'gitlab ci', 'github actions', 'circleci',
        # Other skills
        'git', 'rest api', 'restful', 'graphql', 'microservices', 
        'agile', 'scrum', 'jira', 'linux', 'nginx', 'apache',
        'machine learning', 'ml', 'data science', 'data analyst',
        'html', 'css', 'tailwind', 'bootstrap', 'sass', 'figma',
        'problem solving', 'time management', 'communication',
        # QA & Testing
        'qa', 'quality assurance', 'manual testing', 'automation testing',
        'test case', 'test scenario', 'bug reporting', 'regression testing',
        'functional testing', 'selenium', 'appium', 'katalon', 'postman',
        'cypress', 'playwright', 'junit', 'pytest', 'trello', 'clickup',
        'spreadsheet', 'excel', 'google sheets', 'whimsical', 'flowchart',
    ]
    
    # Indonesian cities
    CITIES = [
        'jakarta', 'bandung', 'surabaya', 'yogyakarta', 'jogja', 'semarang',
        'medan', 'makassar', 'palembang', 'tangerang', 'bekasi', 'depok',
        'bogor', 'malang', 'solo', 'denpasar', 'bali', 'remote'
    ]
    
    def parse(self, content: bytes, filename: str, content_type: str) -> ParsedCV:
        """
        Parse CV file and extract structured data
        """
        # Extract text based on file type
        if content_type == "application/pdf":
            text = self._extract_text_from_pdf(content)
        else:
            text = self._extract_text_from_docx(content)
        
        # Clean text
        text_lower = text.lower()
        
        # Extract data
        parsed = ParsedCV(
            filename=filename,
            name=self._extract_name(text),
            email=self._extract_email(text),
            phone=self._extract_phone(text),
            location=self._extract_location(text_lower),
            education_level=self._extract_education_level(text_lower),
            education_major=self._extract_education_major(text_lower),
            experience_years=self._extract_experience_years(text_lower),
            experience_titles=self._extract_job_titles(text),
            skills=self._extract_skills(text_lower),
            bootcamps=self._extract_bootcamps(text_lower),
            portfolio_urls=self._extract_urls(text),
            github_url=self._extract_github(text),
            linkedin_url=self._extract_linkedin(text),
            raw_text=text[:2000]  # Store first 2000 chars for debugging
        )
        
        return parsed
    
    def _extract_text_from_pdf(self, content: bytes) -> str:
        """Extract text from PDF using PyMuPDF with improved layout handling"""
        text = ""
        try:
            doc = fitz.open(stream=content, filetype="pdf")
            for page in doc:
                # Try dict mode for better layout understanding
                blocks = page.get_text("dict")["blocks"]
                text_blocks = []
                
                for block in blocks:
                    if block.get("type") == 0:  # Text block
                        block_text = ""
                        for line in block.get("lines", []):
                            line_text = ""
                            for span in line.get("spans", []):
                                line_text += span.get("text", "")
                            block_text += line_text + " "
                        
                        # Store with position for sorting
                        x0 = block.get("bbox", [0])[0]
                        y0 = block.get("bbox", [0, 0])[1]
                        text_blocks.append((y0, x0, block_text.strip()))
                
                # Sort by y first (top to bottom), then x (left to right)
                text_blocks.sort(key=lambda b: (b[0] // 50, b[1]))  # Group by ~50px rows
                
                for _, _, block_text in text_blocks:
                    if block_text:
                        text += block_text + "\n"
                
            doc.close()
        except Exception as e:
            print(f"Error extracting PDF: {e}")
            # Fallback to simple extraction
            try:
                doc = fitz.open(stream=content, filetype="pdf")
                for page in doc:
                    text += page.get_text()
                doc.close()
            except:
                pass
        return text
    
    def _extract_text_from_docx(self, content: bytes) -> str:
        """Extract text from DOCX"""
        text = ""
        try:
            doc = Document(BytesIO(content))
            for para in doc.paragraphs:
                text += para.text + "\n"
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
        except Exception as e:
            print(f"Error extracting DOCX: {e}")
        return text
    
    def _extract_email(self, text: str) -> Optional[str]:
        """Extract email address"""
        match = re.search(self.EMAIL_PATTERN, text)
        return match.group(0) if match else None
    
    def _extract_phone(self, text: str) -> Optional[str]:
        """Extract phone number"""
        match = re.search(self.PHONE_PATTERN, text)
        return match.group(0) if match else None
    
    def _extract_name(self, text: str) -> Optional[str]:
        """Extract name (usually first line or after 'Nama:')"""
        lines = text.strip().split('\n')
        for line in lines[:5]:  # Check first 5 lines
            line = line.strip()
            # Skip if line contains email or phone
            if '@' in line or re.search(r'\d{4}', line):
                continue
            # Skip if too short or too long
            if 2 < len(line) < 50:
                # Skip common headers
                if line.lower() not in ['curriculum vitae', 'cv', 'resume', 'profile']:
                    return line
        return None
    
    def _extract_location(self, text_lower: str) -> Optional[str]:
        """Extract location/city"""
        for city in self.CITIES:
            if city in text_lower:
                return city.title()
        return None
    
    def _extract_education_level(self, text_lower: str) -> Optional[str]:
        """Extract highest education level with improved accuracy"""
        # Priority order: Check specific keywords first (more reliable)
        # Check for Bachelor/S1 keywords first (most common)
        s1_keywords = ['bachelor', "bachelor's", 'sarjana', 's.kom', 'skom', 's.t', 's.e']
        for kw in s1_keywords:
            if kw in text_lower:
                return 'S1'
        
        # Check S2 keywords
        s2_keywords = ['master', "master's", 'magister', 'msc', 'mba', 'mm', 'mt', 'mti']
        for kw in s2_keywords:
            if kw in text_lower:
                return 'S2'
        
        # Check S3 keywords
        s3_keywords = ['doktor', 'phd', 'doctoral', 'doctorate', 'doctor of philosophy']
        for kw in s3_keywords:
            if kw in text_lower:
                return 'S3'
        
        # Check D3/D4 keywords
        d3_keywords = ['diploma', 'ahli madya', 'd3', 'd-3']
        for kw in d3_keywords:
            if re.search(rf'\b{re.escape(kw)}\b', text_lower):
                return 'D3'
        
        d4_keywords = ['sarjana terapan', 'd4', 'd-4']
        for kw in d4_keywords:
            if re.search(rf'\b{re.escape(kw)}\b', text_lower):
                return 'D4'
        
        # Check SMA/SMK
        sma_keywords = ['sma', 'smk', 'high school', 'smu', 'slta']
        for kw in sma_keywords:
            if re.search(rf'\b{re.escape(kw)}\b', text_lower):
                return 'SMA'
        
        # Fallback: Check standalone S1/S2/S3 with word boundary (avoid AWS S3)
        # S1 pattern - standalone S1 not followed by bucket/storage/aws terms
        if re.search(r'\bs1\b(?!\s*(?:bucket|storage))', text_lower):
            return 'S1'
        if re.search(r'\bs2\b', text_lower):
            return 'S2'
        # S3 pattern - avoid AWS S3, Amazon S3
        if re.search(r'(?<!aws\s)(?<!amazon\s)\bs3\b(?!\s*(?:bucket|storage|aws|amazon))', text_lower):
            return 'S3'
        
        return None
    
    def _extract_education_major(self, text_lower: str) -> Optional[str]:
        """Extract education major"""
        majors = [
            ('informatika', 'Teknik Informatika'),
            ('ilmu komputer', 'Ilmu Komputer'),
            ('sistem informasi', 'Sistem Informasi'),
            ('teknik komputer', 'Teknik Komputer'),
            ('computer science', 'Computer Science'),
            ('information technology', 'Information Technology'),
            ('software engineering', 'Software Engineering'),
            ('elektro', 'Teknik Elektro'),
            ('matematika', 'Matematika'),
            ('statistik', 'Statistik'),
        ]
        for keyword, name in majors:
            if keyword in text_lower:
                return name
        return None
    
    def _extract_experience_years(self, text_lower: str) -> Optional[int]:
        """Extract years of experience"""
        # Look for patterns like "5 years", "3+ years", "5 tahun"
        patterns = [
            r'(\d+)\+?\s*(?:years?|tahun)\s*(?:of\s*)?(?:experience|pengalaman)?',
            r'(?:experience|pengalaman)\s*(?:of\s*)?(\d+)\+?\s*(?:years?|tahun)',
        ]
        for pattern in patterns:
            match = re.search(pattern, text_lower)
            if match:
                return int(match.group(1))
        return None
    
    def _extract_job_titles(self, text: str) -> List[str]:
        """Extract job titles"""
        titles = []
        title_keywords = [
            'developer', 'engineer', 'programmer', 'analyst', 'manager',
            'lead', 'senior', 'junior', 'architect', 'consultant',
            'specialist', 'administrator', 'designer', 'scientist'
        ]
        text_lower = text.lower()
        for keyword in title_keywords:
            if keyword in text_lower:
                # Find the full title
                pattern = rf'\b\w*\s*{keyword}\b'
                matches = re.findall(pattern, text_lower)
                titles.extend([m.strip().title() for m in matches[:3]])
        return list(set(titles))[:5]
    
    def _extract_skills(self, text_lower: str) -> List[str]:
        """Extract technical skills"""
        found_skills = []
        for skill in self.SKILL_KEYWORDS:
            # Use word boundary for accurate matching
            if re.search(rf'\b{re.escape(skill)}\b', text_lower):
                found_skills.append(skill.upper() if len(skill) <= 3 else skill.title())
        return list(set(found_skills))
    
    def _extract_bootcamps(self, text_lower: str) -> List[str]:
        """Extract bootcamp/training mentions"""
        found = []
        for provider in self.BOOTCAMP_PROVIDERS:
            if provider in text_lower:
                found.append(provider.title())
        return list(set(found))
    
    def _extract_urls(self, text: str) -> List[str]:
        """Extract all URLs"""
        return re.findall(self.WEBSITE_PATTERN, text)
    
    def _extract_github(self, text: str) -> Optional[str]:
        """Extract GitHub URL"""
        match = re.search(self.GITHUB_PATTERN, text, re.IGNORECASE)
        return match.group(0) if match else None
    
    def _extract_linkedin(self, text: str) -> Optional[str]:
        """Extract LinkedIn URL"""
        match = re.search(self.LINKEDIN_PATTERN, text, re.IGNORECASE)
        return match.group(0) if match else None
